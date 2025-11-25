from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# --- Small helpers -----------------------------------------------------------

def set_normal_style(doc, font_name="Calibri", font_size=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    return p


def add_bullet_list(doc, items):
    for txt in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(txt)


def add_contact_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    r = p.add_run(
        "Joakim Pettersson\n"
        "Dalby, Sweden\n"
        "Phone: +46 708 29 99 74\n"
        "Email: joakim.pettersson@additude.se\n"
        "Consultant via ICT Additude AB"
    )
    r.bold = False
    r.font.size = Pt(10)


# --- Cover letter generation -------------------------------------------------

def generate_cover_letter(filename="MAXIV_BioMAX_cover_letter.docx"):
    doc = Document()
    set_normal_style(doc)

    add_contact_block(doc)

    p = doc.add_paragraph()
    p.add_run("To the hiring committee,\nBioMAX – MAX IV Laboratory").bold = False

    heading = doc.add_paragraph()
    run = heading.add_run(
        "Application for temporary Research Engineer position at BioMAX"
    )
    run.bold = True
    run.font.size = Pt(12)

    # --- Body ----------------------------------------------------------------
    body = [
        (
            "I am writing to express my interest in supporting BioMAX as a temporary "
            "Research Engineer during the period you need coverage. I have followed "
            "MAX IV for many years, and when I saw this opening I immediately felt "
            "that my background in instrumentation, motion systems, real-time "
            "measurement and research-adjacent engineering aligns very well with "
            "your needs."
        ),
        (
            "For more than twenty years I have been building, commissioning, "
            "debugging and integrating complex measurement and control systems at "
            "the intersection of electronics, mechanics, vacuum/cryo equipment, "
            "data acquisition and real-time software. Below I highlight a few "
            "experiences that are particularly relevant for BioMAX."
        ),
    ]
    for para in body:
        doc.add_paragraph(para)

    # Motion systems & instrumentation
    add_heading(doc, "Motion systems and instrumentation", level=2)
    add_bullet_list(doc, [
        "Radar sensor calibration motion system at Saab Tank Radar "
        "(PL/M, BASIC, DOS; 1995–1996).",
        "Cryo- and beam-related motion systems for Molecular Beam Epitaxy (MBE) "
        "and E-beam Lithography at Chalmers (Pascal, VMS; 1992–1996), including "
        "a publication on electron-beam scattering in multilayer stacks and its "
        "use for self-aligned e-beam lithography.",
        "Tri-axis motor control for alignment of a pickup system for ElonRoad "
        "electric road charging in heavy vehicles (precision motion under harsh "
        "mechanical and electrical conditions).",
        "Automotive motion systems: drivetrain simulation with drive cycles, "
        "rollover warning systems for trucks, drivetrain fault injection and "
        "gearshift comfort detection (Simulink, MATLAB, C, DOS; Volvo, 1997–2000).",
        "Motor control within fluid-conditioning systems: an energy-recycling "
        "shower PLC (MicroPython, Python, Linux; Orbital Systems) and a "
        "heat-exchanger test system (LabVIEW FPGA/RT, LabVIEW, Python; "
        "Gambro/Baxter).",
    ])

    # Detectors, radiation, ESS
    add_heading(doc, "Detectors, radiation environments and commissioning", level=2)
    add_bullet_list(doc, [
        "Work at the European Spallation Source (ESS), next door to MAX IV, "
        "focused on integration between ESS Python and standard Python "
        "instrument drivers in order to bypass driver backlogs for a one-shot "
        "audit and EMI verification task.",
        "Design of a general signal 'shape-up' guard circuit for ESS real-time "
        "signals, to eliminate reflections at both ends of signal cables even "
        "after EMI issues had been mitigated.",
        "Use of this circuitry during commissioning of radiation monitoring and "
        "beamline alignment equipment.",
    ])

    # Lasers / optics
    add_heading(doc, "Laser and optical instrumentation", level=2)
    add_bullet_list(doc, [
        "Early experience with a LIDAR spectrometer as a student project at IVL "
        "(1991–1992).",
    ])

    # Software / Python
    add_heading(doc, "Software, Python and control", level=2)
    add_bullet_list(doc, [
        "Primary use of Python since 2004 for automation, data handling, "
        "instrument drivers, and analysis scripts.",
        "Extensive experience with embedded C/C++, RTOS-based systems and "
        "FPGA-adjacent integration in several projects.",
    ])

    # Communication & collaboration
    add_heading(doc, "Communication and collaboration with scientists", level=2)
    doc.add_paragraph(
        "I greatly enjoy collaborating with field engineers and scientists from "
        "any discipline. I am comfortable discussing science and technology in "
        "depth, and have repeatedly served as the technical bridge between "
        "hardware teams, software developers and research groups."
    )
    add_bullet_list(doc, [
        "Product Integration Responsible for a complete mobile phone project "
        "(Sony Ericsson C905 camera phone).",
        "Customer Support Engineer at Ericsson Microelectronics in Kista "
        "(2000–2001), focusing on qualification processes for Bluetooth radios "
        "and customer integrations, including conference presentations."
    ])

    # Availability / close
    closing = [
        (
            "I live in Dalby, within easy cycling distance of MAX IV, and I am "
            "available to start on short notice. Whether this temporary position "
            "is due to parental leave or another form of absence, I would be very "
            "happy to cover the full period and contribute immediately to "
            "instrument operation, motion systems, commissioning and user support "
            "at BioMAX."
        ),
        (
            "I would warmly welcome the opportunity to discuss how I can support "
            "your team during this period."
        ),
    ]
    for para in closing:
        doc.add_paragraph(para)

    p = doc.add_paragraph()
    p.add_run("Kind regards,\n").bold = False
    p.add_run("Joakim Pettersson").bold = True

    doc.save(filename)


# --- CV generation -----------------------------------------------------------

def generate_cv(filename="MAXIV_BioMAX_CV.docx"):
    doc = Document()
    set_normal_style(doc)

    # Name as title
    title = doc.add_paragraph()
    r = title.add_run("Joakim Pettersson")
    r.bold = True
    r.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    subtitle = doc.add_paragraph()
    subtitle.add_run("Senior Scientific Software & Electronics Engineer").italic = True

    add_contact_block(doc)

    # Profile
    add_heading(doc, "Profile", level=2)
    doc.add_paragraph(
        "Engineer with more than 20 years of experience in instrumentation, "
        "real-time measurement, motion systems, signal processing and system "
        "integration. Comfortable working in research environments with "
        "detectors, radiation-exposed electronics, beamline-like setups and "
        "complex experimental control systems."
    )

    # Key skills
    add_heading(doc, "Key skills", level=2)
    add_bullet_list(doc, [
        "Instrumentation: electronics, motion systems, cryo/vacuum-related equipment.",
        "Motion systems: radar calibration, MBE and e-beam lithography, "
        "electric-road pickup alignment, drivetrain simulation and PLC-controlled "
        "fluid systems.",
        "Data acquisition and commissioning in demanding environments "
        "(ESS, automotive, medical, radar).",
        "Python (primary language since 2004), C/C++, MATLAB/Simulink, LabVIEW RT/FPGA.",
        "Debugging and hardening of real-time signal chains under EMI and "
        "reflection-prone conditions.",
        "Collaboration with scientists, field engineers and product teams; "
        "user support during experiments.",
    ])

    # Relevant experience – condensed
    add_heading(doc, "Selected relevant experience", level=2)

    exp = doc.add_paragraph()
    exp.add_run("European Spallation Source (ESS) – Electronics / DAQ engineering\n").bold = True
    doc.add_paragraph(
        "Integration of ESS Python with standard Python instrument drivers for "
        "a focused audit and EMI verification task. Design of general signal "
        "conditioning circuitry to remove reflections on real-time signals, used "
        "during commissioning of radiation monitoring and beamline alignment "
        "equipment."
    )

    exp = doc.add_paragraph()
    exp.add_run("Chalmers University – MBE and E-beam lithography motion systems\n").bold = True
    doc.add_paragraph(
        "Motion systems involving cryo equipment, vacuum systems and beam control "
        "in Molecular Beam Epitaxy and E-beam Lithography setups (Pascal, VMS). "
        "Published work on electron-beam scattering in material stacks and its "
        "use in self-aligned e-beam lithography."
    )

    exp = doc.add_paragraph()
    exp.add_run("ElonRoad – Electric road systems for heavy vehicles\n").bold = True
    doc.add_paragraph(
        "Tri-axis motor control for precise alignment of a pickup system that "
        "draws power from the electrified road in heavy vehicles. Real-time "
        "control, sensor integration and robust operation in harsh conditions."
    )

    exp = doc.add_paragraph()
    exp.add_run("Saab Tank Radar – Radar sensor calibration systems\n").bold = True
    doc.add_paragraph(
        "Development of motion systems for radar sensor calibration (PL/M, BASIC, "
        "DOS). Involvement with RF-adjacent measurement setups and instrumentation."
    )

    exp = doc.add_paragraph()
    exp.add_run("Volvo – Automotive motion and drivetrain systems\n").bold = True
    doc.add_paragraph(
        "Vehicle drivetrain simulation with drive cycles, rollover warning for "
        "trucks, drivetrain fault-injection and gear-shift comfort detection "
        "(Simulink,
