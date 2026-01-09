# cv.py
"""
Class-based CV & cover-letter generator for Joakim Pettersson.

Content corresponds to the compressed 2025-11-11 CV, with ONE deliberate change:

- The first PROFILE bullet ("Embedded software developer with 14+ years...")
  is removed to keep the base CV less branch-specific.

Styling is based on the original procedural cv.py:
- A4 page, 25/20 mm margins
- Calibri 10 pt
- Accent blue section headings
- Bold dark-blue hyperlinks (no underline)
- Hanging indents for Technology and artifact lines
"""

from __future__ import annotations
from dataclasses import dataclass
from email.contentmanager import maintype

from numpy.lib.arraysetops import union1d
from numpy.lib.recfunctions import structured_to_unstructured
from regex import compile
from regex._regex_core import error as RegexError
from os.path import splitext

from pymupdf import Document as PdfDocument
import docx
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from traits.api import HasTraits, File, List, Str, Tuple
from traits.trait_notifiers import push_exception_handler
from numpy import array, argwhere, sort, empty
from traitsui.api import View, Item, HGroup, Group
from traitsui.editors.api import ListEditor, TextEditor

from ui.str_cell_editor import FlowStrEditor
from ui.flow_list_str_editor import FlowListStrEditor

push_exception_handler(reraise_exceptions=True)
Text = Str(editor=FlowStrEditor())
ListText = List(Str, editor=FlowListStrEditor())
Line = Str(editor=FlowStrEditor(max_lines=1))
ListLine = List(Str, editor=FlowListStrEditor(max_lines=1))


# ---- Branding / colors ----
ACCENT_BLUE = RGBColor(0x00, 0x66, 0xB3)  # Additude headings
LINK_BLUE_HEX = "004A99"                  # Bold dark blue for hyperlinks (no underline)

# =====================================================================
# StyledDocument – wrapper around python-docx.Document
# =====================================================================


class StyledDocument:
    """Wrapper around docx.Document with shared layout and style helpers."""

    def __init__(self) -> None:
        self.doc: DocxDocument = docx.Document()
        self._set_page()
        self._set_base_style()

    # --- Page + base styles ----------------------------------------------

    def _set_page(self) -> None:
        s = self.doc.sections[0]
        # A4
        s.page_width, s.page_height = Mm(210), Mm(297)
        # Margins
        s.left_margin = s.right_margin = Mm(25)
        s.top_margin = Mm(20)
        s.bottom_margin = Mm(20)

    def _set_base_style(self) -> None:
        st = self.doc.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10)

    # --- Section & paragraph helpers -------------------------------------

    def add_section_heading(
        self,
        text: str,
        *,
        level: int = 1,
        space_before: int = 18,
        space_after: int = 6,
    ):
        """
        Add a colored heading using Word's heading styles.

        `text` is uppercased for visual consistency.
        """
        h = self.doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = ACCENT_BLUE

        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        return h

    def add_company_role_title(self, company, url, role):
        """
        Example:
            add_company_role_title(
                doc,
                "Elonroad",
                "https://www.elonroad.com",
                " (2025) – Software Developer, Lund",
            )
        """
        p = self.add_section_heading("", level=3, space_before=8, space_after=0)

        # Linked, bold company name
        self.add_hyperlink(p, company, url, bold=False)

        # Rest of the line (years, role, location)
        p.add_run(" " + role)

        return p

    def add_para(self, text: str, space_before: int = 0, space_after: int = 2, left_indent: float = 0.5, right_indent: float = 0.):
        """
        Normal paragraph with tight rhythm (0 before, 2 after).
        """
        p = self.doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = Cm(left_indent)
        p.paragraph_format.right_indent = Cm(right_indent)
        return p

    def add_bullet(self, text: str):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5 + 0.63)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        p.add_run(text)
        return p

    # --- Hyperlink + artifact + technology -------------------------------

    def add_hyperlink(self, paragraph, text: str, url: str,
                      bold: bool = True, color_hex: str = LINK_BLUE_HEX, underline: bool = False):
        """
        Add a clickable hyperlink to `paragraph` with custom styling.
        """

        if not url:
            return paragraph.add_run(text)

        part = paragraph.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Color
        color = OxmlElement("w:color")
        color.set(qn("w:val"), color_hex)
        rPr.append(color)

        # Bold?
        if bold:
            rPr.append(OxmlElement("w:b"))

        # Underline?
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single" if underline else "none")
        rPr.append(u)

        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)

        link.append(run)
        paragraph._p.append(link)
        return paragraph

    def add_artifact(self, title: str, url: str):
        """
        Add a hanging-indent artifact line:

            → Title (hyperlink)

        with left indent and negative first-line indent.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.add_run("→ ")
        self.add_hyperlink(p, title, url)
        return p

    def add_tech(self, *items: str):
        """
        Add a hanging-indent Technology line:

            Technology: <comma-separated items>   (italic)
        """
        if not items:
            return None
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run("Technology: ")
        r1.bold = True
        r1.italic = True
        r2 = p.add_run(", ".join(items))
        r2.italic = True
        return p

    # --- Save ------------------------------------------------------------

    def save(self, filename: str) -> None:
        self.doc.save(filename)


# =====================================================================
# BaseCV – compressed 2025-11-11 CV (with 1 profile bullet removed)
# =====================================================================


def joined(items: list[str], sep=", ", final_sep=" & "):
    n = len(items)
    if n == 0:
        return ""
    if n == 1:
        return items[0]
    return sep.join(items[:-1]) + final_sep + items[-1]


def flat_text_list_editor(rows: int = 6, style="custom") -> ListEditor:
    """
    A compact, per-item text editor list factory (no listbox).
    - style='custom' lays out editors vertically
    - editor=TextEditor(...) gives plain text fields per row
    - rows controls vertical size (adds scroll if needed)
    """
    return ListEditor(
            style=style,
            editor=TextEditor(auto_set=False, enter_set=True),
            rows=rows,              # target visible rows before scrolling
        )


class BaseCV(StyledDocument, HasTraits):
    """Base CV corresponding to the compressed 2025-11-11 version."""

    # ----------------- Header --------------------------------------------

    name = Str
    role = Str
    level = Str
    specialities = ListLine
    industries = ListLine
    email = Str
    phone = Str
    linkedin_profile = Str

    def __init__(self, **kwargs) -> None:
        StyledDocument.__init__(self)
        HasTraits.__init__(self, **kwargs)

    def __iadd__(self, other):
        """Merge other into self"""
        old_size = getattr(self, "size", None) or 1
        new_size = getattr(other, "size", None) or 1
        self.size = max(old_size, new_size)
        for attr, handler in self.traits().items():
            if attr.endswith('_positions') or attr.endswith('_index'):
                continue

            if hasattr(other, attr):
                old = getattr(self, attr)
                new = getattr(other, attr)
                if not old and not new:
                    continue

                info = handler.full_info(self, attr, None)
                if info == 'a string':
                    old_positions = getattr(self, attr + '_positions', array([0, old_size]))
                    new_positions = getattr(other, attr + '_positions', array([0, new_size]))
                    positions, value = (new_positions, new) if new else (old_positions, old)
                    setattr(self, attr, value)
                    setattr(self, attr + '_positions', positions)
                elif info == 'a list of items which are a string':
                    old_positions = list(getattr(self, attr + '_positions', [])) or list(range(len(old)))
                    new_positions = list(getattr(other, attr + '_positions', [])) or list(range(len(new)))

                    # Keep only unique values and in the same order
                    position_values = []
                    values = []
                    for position, value in sorted(zip(old_positions + new_positions, old + new)):
                        if not value in values:
                            values.append(value)
                            position_values.append((position, value))

                    positions, values = zip(*position_values)
                    setattr(self, attr, list(values))
                    setattr(self, attr + '_positions', array(positions + (max(old_size, new_size),)))

    def add_header(self) -> None:
        # Name + subtitle in a tight block
        p = self.doc.add_heading()
        p.paragraph_format.space_after = Pt(0)
        p.add_run(self.name).font.size = Pt(18)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{self.level} {joined(self.specialities)} {self.role} – {joined(self.industries)}"
                  ).font.size = Pt(12)

        # Contact line
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(f"📧 {self.email}  📱 {self.phone}  🔗 ")
        # Hyperlink for LinkedIn
        self.add_hyperlink(p, self.linkedin_profile.rstrip('/').lstrip('https://www.'), self.linkedin_profile)

    # ----------------- Profile -------------------------------------------

    profile = List(Text)

    def add_profile(self) -> None:
        self.add_section_heading("PROFILE")
        for bullet in self.profile:
            self.add_bullet(bullet)

    # ----------------- Core competence -----------------------------------

    core_competence = List(Tuple(Str, ListLine))

    def add_core_competence(self) -> None:
        self.add_section_heading("CORE COMPETENCE")
        for category, items in self.core_competence:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{category}: ")
            r.bold = True
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.add_run(" • ".join(items))

    # ----------------- Experience ----------------------------------------

    Experience = Tuple(
        Tuple(Str, Str, Str),
        List(Text),
        List(Text),
        List(Str),
        List(Tuple(Str, Str)))
    experience = List(Tuple(Str, List(Experience)))

    LINK = compile(r"\[\**([^\*\]]*)\**\]\(([^)]*)\)")  # [text](url)

    def add_experience(self) -> None:
        for heading, experience in self.experience:
            self.add_section_heading(heading.upper())
            for (company, company_url, role), description, bullets, technologies, artifacts in experience:
                self.add_company_role_title(company, company_url, role)
                for text_method, texts in [(self.add_para, description), (self.add_bullet, bullets)]:
                    for text in texts:
                        i = 0
                        for m in self.LINK.finditer(text):
                            if i:
                                p.add_run(text[i:m.start()])
                            else:
                                p = text_method(text[0:m.start()])

                            self.add_hyperlink(p, *m.groups())
                            i = m.end()

                        if i:
                            p.add_run(text[i:])
                        else:
                            text_method(text)

                self.add_tech(*technologies)
                for artifact, url in artifacts:
                    self.add_artifact(artifact, url)

    # ----------------- Working approach & personal ------------------------

    environment_approaches = List(Tuple(Str, List(Text)))

    def add_working_approach(self) -> None:
        for environment, approaches in self.environment_approaches:
            self.add_section_heading(environment.upper())
            for approach in approaches:
                self.add_bullet(approach)

    # ----------------- Top-level build -----------------------------------

    def build_cv(self, filename: str) -> None:
        self.add_header()
        self.add_profile()
        self.add_core_competence()
        self.add_experience()
        self.add_working_approach()
        self.save(filename)
        print(self.__class__.__name__, "saved to", filename)


# =====================================================================
# BaseCoverLetter – neutral, extendable
# =====================================================================

class BaseCoverLetter(BaseCV):
    """Generic, neutral cover letter structure (content intentionally broad)."""
    role = Str
    organization = Str
    to_ask = Str
    to = Str
    location = Str
    affiliation = Str
    work = Text
    motivation = Text
    arguments = List(Text)
    hook = Text

    def add_field(self, paragraph, style_name: str, text: str, eols: int = 1):
        # 1) create a character style (always new as you wanted)
        style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)

        # 2) add run and *force* rStyle onto the run's rPr so Pandoc keeps it
        run = paragraph.add_run(text)
        r = run._r
        rPr = r.get_or_add_rPr()
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), style_name)
        # ensure rStyle exists even if no visible font deltas
        rPr.insert(0, rStyle)

        # 3) wrap with SDT tagged the same
        sdt = OxmlElement('w:sdt')
        pr = OxmlElement('w:sdtPr')
        tag = OxmlElement('w:tag');
        tag.set(qn('w:val'), style_name)
        pr.append(tag)
        pr.append(OxmlElement('w:text'))
        sdt.append(pr)
        content = OxmlElement('w:sdtContent')
        r.addprevious(sdt)
        content.append(r)
        sdt.append(content)
        for i in range(eols):
            run.add_break(WD_BREAK.LINE)

        return run

    def add_contact_block(self) -> None:
        p = self.doc.add_paragraph()

        r = self.add_field(p, "name", self.name)
        r = self.add_field(p, "location", self.location)
        r = self.add_field(p, "phone", self.phone)
        r = self.add_field(p, "email", self.email)
        r = self.add_field(p, "engagement", self.affiliation, 2)

        r = p.add_run(self.to_ask)
        r = self.add_field(p, "to", f"{self.to},")
        r = self.add_field(p, "organization", self.organization)

    def add_opening(self) -> None:
        p = self.add_section_heading("Application for ")
        self.add_field(p, "role", self.role, 0)

    def add_intro(self) -> None:
        self.doc.add_paragraph(
            f"I am writing to express my interest in supporting {self.organization} as {self.role} "
            f"working with {self.work}. With {self.motivation}, I believe I can contribute from day one:"
        )

    def add_body(self) -> None:
        for argument in self.arguments:
            self.add_para(argument, 18, 18, 1.5, 1.5)

    def add_closing(self) -> None:
        self.doc.add_paragraph(self.hook)
        p = self.doc.add_paragraph()
        p.add_run("Kind regards,\n")
        p.add_run(self.name).bold = True

    def build_cover(self, filename: str) -> None:
        self.__init__()  # Clear doc
        self.add_contact_block()
        self.add_opening()
        self.add_intro()
        self.add_body()
        self.add_closing()
        self.save(filename)
        print(self.__class__.__name__, "saved to", filename)


class MatchError(RuntimeError):
    def __init__(self, hint, pattern_name, content_name, pattern, text, start, end,
                 match = None, troubleshooting=False) -> None:
        hint += (
            f"\nRemaining to match after last <?P<{match.lastgroup}>...):\n{text[match.end():]}\n" if match.end() != end else
            f"\nTrouble with last <?P<{match.lastgroup}>...) match: {repr(text[slice(*match.regs[-1])])}\n"
        ) if match and troubleshooting else (
            "Remember to add back \\Z at end of this pattern after troubleshooting.\n" if troubleshooting else
            "Temporarily remove \\Z at end of pattern to troubleshoot!\n")
        raise SyntaxError((f"{content_name} text does not match Python https://regex101.com/\n"
                           "FLAVOR: Python\n"
                           f"REGULAR EXPRESSION:\n{pattern}\n"
                           f"TEST STRING:\n{text[start:end]}\n"
                           "- All (indented) patterns above must match with the file text.\n"
                           f"{hint}"
                           f"Correct {pattern_name} pattern or {content_name} content until they match.").replace(
            '(?>', '(?:').replace(
            r'\G', r'\A').replace(
            r'\h', r'[^\S\n]').replace(
            '📧', r'\U0001F4E7').replace(
            '📱', r'\U0001F4F1').replace(
            '🔗', r'\U0001F517'))


def match_or_stop(pattern_name, content_name,
                  pattern, text, start=0, end=None):
    """Match or fail with hint

    Refers to https://regex101.com for troubleshooting
    """

    end = end or len(text)
    hint = ""
    try:
        match = compile(pattern).match(text, start, end, timeout=1.0)
    except RegexError as e:
        match = None
        hint = '\n'.join(pattern.split('\n')[:e.lineno])
        hint += f"\n{' ' * (e.colno - 1)}^--- {e.msg} \n"
    except TimeoutError:
        match = None
        hint = "Too deep recursion in regex, split in smaller pattern chunks!\n"

    troubleshooting = pattern[:4] != "(?s)" and r'\Z' not in pattern[-3:]
    if not match or match.start() != start or match.end() != end or troubleshooting:
        raise MatchError(hint, pattern_name, content_name, pattern, text, start, end, match, troubleshooting)

    return match


CR = "(?: <br>)"
EOL = rf"{CR}? \n"
SEP = rf"(?: \h* \n | \h* <br> \n | \h+)"
CHAR = r"[^<\n#]"
TEXT = f"{CHAR}*"
PHONE = "[+ ,0-9]*"
EMAIL = "[A-Za-z0-9_.+-]+@[A-Za-z0-9.-]*"
COMMA_EOL = f"\s* , \s* {EOL}"

def spaceout(text: str) -> str:
    return text.replace(" ", " \h+ ")

def choice(name, *alternatives):
    return f"\h* (?: (?P<{name}> {' | '.join(map(spaceout, alternatives))}) :? \h*)"

def ask(name, pattern, *alternatives):
    return f"(?: {choice(f'{name}_ask', *alternatives)}? (?P<{name}> {pattern}) {SEP})"

ROLE = choice('role', 'Engineer', 'Programmer', 'Developer', 'Designer', 'Manager')
LETTER = rf"""
(?P<add_contact_block> 
  {ask('name', TEXT, 'Full name', 'Name')}?
  {ask('location', TEXT, 'Home address', 'Address')}?
  {ask('phone', PHONE, 'Phone', 'Mobile')}?
  {ask('email', EMAIL, 'Email')}?
  {ask('affiliation', TEXT, 'Affiliation')}?
  {EOL}*
  (?: \h* (?: (?P<to_ask> To :? \h*))? (?P<to> (?> (?! {COMMA_EOL}) [^<\n#])*) {COMMA_EOL})?
  {ask('organization', TEXT, 'Organization')}?
)
  \s*
(?P<add_opening>
  (^ \#+)? {ask('role', TEXT, 'Application for ')}
)
  \s*
(?P<add_intro>
  (?:
    working \h+ with \h+ (?P<work> (?> [^\.]* ) ) \. |
    With \h+ (?P<motivation> (?> (?! , \h+ (?: I | we) \h) . )* ) , \h+ (?P<group> I | we) \h+ |
    (?P=organization) | (?P=role) |
    {CHAR}
  )+ {EOL}
)
(?P<add_body>
  (?: \s* ^ > \h+ (?P<arguments> (?> (?! \n\n ) . )* ) \n\n)*
)
(?P<add_closing> 
  \s* (?P<hook> (?> (?! Sincerely | Kind | Greetings | \** (?P=name) | ^ ---+ \n ) {TEXT} )* ) \n\n
  \s* (?: (?! \** (?P=name) | ^ ---+ \n) {TEXT} {EOL})*
  \s* \** (?P=name) \**
)
"""

CV = rf"""
(?P<add_header>
    \#* \h* (?P=name) \n\n
    
    \h* (?P<level> Senior | Junior)
    \h* (?: (?P<specialities> (?: (?! (?P=role)) [^,&<\n])*) [, &]+)*
    (?P=role)?
    \h [-–—] \h
    (?: (?P<industries> (?: (?! \h* [,&\n]) .)+) \h* [,&]? \h*)* \n\n
    
    \h* (?: Email | 📧)? \h* (?P=email)
    \h* (?: Phone| Mobile | 📱)? \h* (?P=phone)
    \h* (?: LinkedIn | 🔗)? :? \h*
    \[ \** (?: linkedin.com/in/)? (?P<linkedin_name> [^/*\]<\n]*) /? \** \]
    \( (?P<linkedin_profile> (?: https?://www.)? linkedin.com/in/(?P=linkedin_name) /?)? \) \n\n
)
(?P<add_profile>
    \#* \h PROFILE\n\n
    (?: - \h (?P<profile> (?> (?! \n\n ) . )* ) \n\n)*
)
(?P<add_core_competence>
    \#* \h CORE \h COMPETENCE\n\n
    (?:
        \** (?P<_categories> [^:]*) : \** \n\n
        > \h (?: (?: \h•\h)? (?P<_items> (?> (?! \h•\h | \n\n ) . )* ) )* \n\n
    )*
)
(?P<add_experience>
    (?:
        \s* ^ \# \h (?! MENTORSHIP | COLLABORATION | PERSONAL) (?P<_headings> [^\n]*) \n\n
        (?:
            \s* ^ \#\#\# \h \[? (?P<_companies> [^\](]*) \]?
            (?: \( (?P<_company_urls> https?:// [^\)]*) \))? \h
            (?P<_roles> [^\n]*) \n\n
            (?: > \h (?! \** Technology: \** \h | → \h \[) (?P<_descriptions> (?> (?! \n\n ) . )* ) \n\n)*
            (?: - \h (?P<_bullets> (?> (?! \n\n ) . )* ) \n\n)*
            (?: > \h \** Technology: \** \h (?: (?: ,\h)? (?P<_technologies> [^,\*]*))* \*\n\n)?
            (?: > \h → \h \[ \** (?P<_artifacts> [^\*]*) \** \] \( (?P<_artifact_urls> https?:// [^\)]*) \)\n\n)*
        )*
    )*
)
(?P<add_working_approach>
    (?:
        \s* ^ \# \h (?P<_environments> [^\n]*) \n\n
        (?: - \h (?P<_approaches> (?> (?! \n\n ) . )* ) \n\n)*
    )*
)
"""

FILE = rf"""(?mxs)
  \A
(?P<letter> {LETTER})
(?:
  \n\n ---+ \n\n
)
(?P<cv> {CV})
  \Z
"""

INDENTED_SEPARATOR = r"""(?xs) # (?P<{name}[__{value}]>{pattern}) | .*
    \G (?P<indent> [\ \t]* ) (
        \( \? P <
        (?P<name__value>
            (?P<name> (?> (?! __ | >) . )+ )
            (?> __ (?P<value> [^>]* ))?
        )
        >
        (?P<pattern> (?> (?! \) \Z ) . )+ )
        \)
    |   \( \? [>:] (?P<first_of_alternatives> [^|]+ ) .*
    |   .*
    )\Z"""


@dataclass
class Separator:
    indent: str
    section_name: str
    name: str = None
    value: str = None
    pattern: str = None
    capture: str = None


def variablify(s):
    return s.strip().lower().replace(*' _').replace(*'åa').replace(*'äa').replace(*'öo')

def get_separator(indented_separator) -> Separator:
    indent, capture, section_name, name, value, pattern, first_of_alternatives = match_or_stop(
        'INDENTED_SEPARATOR', 'indented_separator',
         INDENTED_SEPARATOR, indented_separator).groups()
    section_name = section_name or first_of_alternatives or variablify(capture)
    section_name = section_name.replace(*'_ ').title().replace(' ', "")  # Unique CamelCase name
    if value:
        value = value.replace(*'_ ')

    return Separator(indent, section_name, name, value, pattern or capture, capture)

DESCRIPTIONS_AND_BULLETS = ( # (m)
    r"(?>(?P<_descriptions>(?>^(?!• ).*\n)*?^(?!• ).*?\.[^\S\n]*)(?>\n+|(?=^• )|\Z))*"
    r"(?>^• (?P<_bullets>.*(?>\n(?!• ).*)*)(?>\n+|\Z))*")

HERMES_PDF_FILE_CHUNKS = rf"""\A(?P<name>.+)
(?>(?P<level>Senior|Junior) )?(?P<specialities> ?(?:(?!\w+$)[\w\-])+)* ?(?P<role>\w+)
\Z

(?>MOTIVATION|MOTIVERING)

\G(?>(?P<arguments>(?>(?>(?!\.\h?\n).)+\n)*.*)\n)*\Z

(?P<_categories__Skills>(?>SKILLS|FÄRDIGHETER))

\G(?>(?:,[\h\n])?(?P<_items>[^,]*))*
?\Z

(?P<_categories__Key_expertise>(?>PROFESSIONAL EXPERTISE|PROFESSIONELL EXPERTIS))

\G(?>(?>[•,][\h\n])?(?P<_items>[^•,]*))*
?\Z

(?P<_headings__EDUCATION>(?>EDUCATIONS|UTBILDNINGAR))

(?m)\G(?P<_companies>.*)
(?P<_roles>.*
\d\d\d\d-\d\d-\d\d - \d\d\d\d-\d\d-\d\d)
{DESCRIPTIONS_AND_BULLETS}\Z

(?P<_headings__EXPERIENCE>(?>EXPERIENCES|ERFARENHETER))

\G\Z

    (?P<_companies__>(?=\d\d\d\d-\d\d-\d\d\n))

\G
?(?P<_roles>\d\d\d\d-\d\d-\d\d)
(?P<_roles___2DH_>UNTIL)
(?P<_roles>\d\d\d\d-\d\d-\d\d)
(?P<_roles___2DH_>ROLE|ROLL)
(?P<_roles>(?:(?>.|
))+)\Z

    (?>COMPANY|FÖRETAG)

(?m)\G(?P<_companies>.+(?:\s*[^\W\d_]+(?:\s+[^\W\d_]+){{0,2}})?)
{DESCRIPTIONS_AND_BULLETS}\Z

    (?>Technologies applied|Teknik som tillämpas)

\G(?>(?>,[ \n])?(?P<_technologies>[^,]*))*\Z

(?P<_categories__Conversations>(?>LANGUAGES|SPRÅK))

\G(?>(?P<_items>.+
.+)
)+.+
\Z"""


hex_chars_regex = compile(r"((?>[0-9a-fA-F]{2})+)(H)")
def name_value(s):
    """Use _ and hexadecimal numbers to insert special characters

    >>> name_value('Special3AH_2AH_2BH_2C2D2E2FH_')
    'Special: * + ,-./ '
    """
    replace = [[s[0] for s in m.allspans()[1:]] for m in hex_chars_regex.finditer(s)]
    for (start, hex_end), (hex_end, end) in reversed(replace):
        s = s[:start] + bytes.fromhex(s[start:hex_end]).decode() + s[end:]

    return s.replace(*'_ ')


def get_key(item):
    """Return first str in item"""
    if isinstance(item, str):
        return item

    return get_key(item[0])


class Proposal(BaseCoverLetter):
    """Handle an assignment proposal file

    Regenerates the same build_cv() and build_cover() Word documents that MarkdownBuilder used for the proposal file.
    """
    file = File()
    size = int
    _categories = _items = _headings = \
        _companies = _company_urls = _roles = _descriptions = _bullets = _technologies = _artifacts = _artifact_urls = \
        _environments = _approaches = List(Str)

    # Assembly information for deeper traits from shallower ones, arranged by position in a proposal markdown file
    structures = (
        ('core_competence', [
            ('_categories', [
                '_items'])]),
        ('experience', [
            ('_headings', [(
                ('_companies', '_company_urls', '_roles'),
                ['_descriptions'],
                ['_bullets'],
                ['_technologies'],
                [('_artifacts', '_artifact_urls')])])]),
        ('environment_approaches', [
            ('_environments', [
                '_approaches'])]),
    )

    # For the edit_traits() method:
    traits_view = View
    file_positions = [-1]  # Every pubic trait needs _positions so that we can sort them all into traits_view.

    def _file_changed(self):
        """Import text from a file"""
        if hasattr(self, 'text'):
            # Create a proposal with just the file and then merge it into the existing one
            self += Proposal(file=self.file)
            return

        ext = splitext(self.file)[1]
        if ext == '.md':
            self.pattern_name = 'FILE'
            self.pattern = FILE
            self.text = open(self.file, encoding="utf-8").read()
        elif ext == '.pdf':
            self.pattern_name = 'HERMES_PDF_FILE_CHUNKS'
            self.pattern = HERMES_PDF_FILE_CHUNKS
            self.text = "\n".join(page.get_text() for page in PdfDocument(self.file))
        else:
            raise NotImplementedError(ext)

        # Split pattern into smaller chunks
        pattern_chunks = self.pattern.split("\n\n")
        section_patterns = pattern_chunks[::2]
        indented_separators = pattern_chunks[1::2]

        # Build a structured parser based on how separators are indented in the pattern
        separators =  list(map(get_separator, indented_separators))
        main_pattern = r"(?m)\A(?P<heading>(?>"
        stops = []
        for separator in separators:
            indentation = len(separator.indent)
            main_pattern += rf"(?!{separator.pattern}"
            stopped = False
            while indentation < len(stops):
                stop = stops.pop()
                if stop:
                    main_pattern += rf"|{stop}|\n\n).*\n)*)*)*"
                    stopped = True

            if not stopped:
                main_pattern += r").*\n)*)"

            if indentation > len(stops):
                assert separator.pattern
                stops += [""] * (indentation - len(stops) - 1) + [separator.pattern]
                main_pattern += "(?>"

            main_pattern += f"{separator.capture}\n*(?P<{separator.section_name}>(?>"

        main_pattern += r".*\n)*)\Z"

        # Import traits from separators
        self.size = len(self.text)
        self.prepositions = {}
        structure = self.parse(main_pattern)

        # Import traits from chunks between separators
        for i, (separator, section_pattern) in enumerate(zip(
                [Separator("", 'heading')] + separators, section_patterns)):
            for (start, end) in structure.spans(separator.section_name):
                self.parse(section_pattern, start, end)

        # Assemble deeper-structured traits from simpler ones
        for attr, structure in self.structures:
            values, positions = self.structured(structure)
            setattr(self, attr, values)
            setattr(self, attr + '_positions', positions)
            setattr(self, attr + '_index', 0)

    def parse(self, chunk_pattern, start=0, end=None):
        """Extract traits and _positions from self.text[start:end] using named groups in chunk_pattern

        Any '__' in a group name is used split it into a trait name and a tail for the value to be used instead.
        In that case, the trait receives tail.replace(*'_ ') at all matched positions instead of the matched texts.
        If that tail is empty (from a pure look-ahead pattern with nothing matched), the trait gets the next
        matched value immediately after that pre-position and that next position is simply dropped.
        """
        match = match_or_stop(self.pattern_name, self.file,
                              chunk_pattern, self.text, start, end)
        group_names = list(match.groupdict())
        coercions = [(n, n.split('__')) for n in group_names if '__' in n]
        actuals = [n for n in group_names if '__' not in n]
        coercers = [c[0] for c in coercions]
        names = [c[1][0] for c in coercions]
        coerced = [name_value(c[1][1]) for c in coercions]
        coercer_names = dict(zip(coercers, names))
        coercer_values = dict(zip(coercers, coerced))
        for group_name in coercers + actuals:
            attr = coercer_names.get(group_name, group_name)
            trait = self.trait(attr)
            if trait:
                values = match.captures(group_name)
                if not values:
                    continue

                old_values = getattr(self, attr)
                old_positions = list(getattr(self, attr + '_positions', [])[:len(old_values)])

                positions = match.starts(group_name)
                if attr != group_name:
                    value = coercer_values[group_name]
                    if not value:
                        # Pop-sorted positions for later actual values
                        self.prepositions[attr] = list(reversed(self.prepositions.get(attr, []) + positions))

                    values = [value] * len(positions)
                else:
                    if attr in self.prepositions:
                        # Replace dummy values at pre-positions with values we now have
                        prepositions = self.prepositions[attr]
                        consumed_position_indices = []
                        for i, position in enumerate(positions):
                            if position < prepositions[-1]:
                                continue

                            old_values[old_positions.index(prepositions.pop())] = values[i]
                            consumed_position_indices.append(i)

                        for i in reversed(consumed_position_indices):
                            del positions[i]
                            del values[i]

                        if not values:
                            continue

                setattr(self, attr + '_index', 0)
                info = trait.handler.full_info(self, attr, None)
                if info == 'a string':
                    assert len(positions) == 1, positions
                    setattr(self, attr, values[0] or "")
                    setattr(self, attr + '_positions', array([positions[0], self.size]))
                elif info == 'a list of items which are a string':
                    old_positions_list = list(old_positions[:len(old_values)])
                    positions_list, values = zip(*sorted(zip(old_positions_list + list(positions), old_values + values)))
                    setattr(self, attr, list(values))
                    setattr(self, attr + '_positions', array(positions_list + (self.size,)))

        return match

    def structured(self, structure, span=None, attr_index=0):
        """Return a str with one value, a tuple or a list; and the _positions span used

        In case of a first str in the structure, the span is updated to not include the previous and next position.
        """
        span = span or [0, self.size]
        if isinstance(structure, str):
            attr = structure
            values = getattr(self, attr)
            if not values:
                return '', span

            if attr_index:
                return "".join(self.unstructured(attr, span)), span

            positions = getattr(self, attr + '_positions')
            index = getattr(self, attr + '_index')
            if index >= len(values) or positions[index] >= span[1]:
                return '', span

            span[0], span[1] = positions[index:index + 2]
            setattr(self, attr + '_index', index + 1)
            return values[index].replace(*'\n ').replace('  ', ''), span

        if isinstance(structure, tuple):
            items = []
            for branch_index, branch in enumerate(structure):
                item, span = self.structured(branch, span, branch_index)
                assert branch_index or item
                items.append(item)

            return tuple(items), span

        # List
        item = structure[0]
        if isinstance(item, str):
            return self.unstructured(item, span), span

        forest = []
        first = True
        while True:
            try:
                tree, span_here = self.structured(item, span[:])
                forest.append(tree)
                if first:
                    span[0] = span_here[0]
            except AssertionError:
                return forest, span

    def unstructured(self, attr, span=None):
        """Grab all non-empty values within a position span after having removed \n"""
        values = getattr(self, attr)
        if not values:
            return values

        span = span or [0, self.size]
        positions = getattr(self, attr + '_positions', [])
        index = getattr(self, attr + '_index')
        next_structure_index = index + argwhere(positions[index:] >= span[1])[0][0]
        setattr(self, attr + '_index', next_structure_index)
        values = values[index:next_structure_index]
        for i, value in reversed(list(enumerate(values))):
            value = value.replace(*'\n ').replace('  ', '')
            if value:
                values[i] = value
            else:
                del values[i]

        return values

    def __iadd__(self, other, structure=None, old=None, new=None):
        if not structure:
            super().__iadd__(other)
            old = [getattr(self, attr, []) for attr, structure in self.structures]
            new = [getattr(other, attr, []) for attr, structure in self.structures]
            for (attr, structure), old, new in zip(self.structures, old, new):
                setattr(self, attr, self.__iadd__(other, structure, old, new))

            return

        if isinstance(structure, str):
            return new or old or ""

        if isinstance(structure, tuple):
            return tuple([self.__iadd__(other, branch, old, new) for branch, old, new in zip(structure, old, new)])

        # List
        tree = structure[0]
        if isinstance(tree, str):
            # Keep only unique values and in the old order, appending any unique new in the new order
            values = []
            for value in old + new:
                if not value in values:
                    values.append(value)

            return values

        # Use first str in each old and new data value as a key, and merge everything else where it was, or append
        items = old[:]
        keys = list(map(get_key, items))
        for key, new_item in zip(map(get_key, new), new):
            if key in keys:
                i = keys.index(key)
                items[i] = self.__iadd__(other, tree, items[i], new_item)
            else:
                items.append(new_item)

        return items

    def view(self):
        """Make a view for all public traits sorted by _positions[0]"""
        pos_attrs = sorted([(getattr(self, attr + '_positions', array([self.size]))[0], attr)
                            for attr in self.traits().keys() if self.is_view_item(attr)])
        attrs = [name for _, name in pos_attrs]
        items = [Item(name, springy=False) for name in attrs]
        splits = [attrs.index(new_group) for new_group in ('level', 'experience', 'environment_approaches')]
        group_indices = [(start, end) for start, end in zip([0] + splits, splits + [len(attrs)])]
        groups = [Group(*items[start:end]) for start, end in group_indices]
        return View(HGroup(*groups),
                                title="Proposal", resizable=True, buttons=["OK"],
                                x=0, y=0, width=1.0, height=1.0)

    def is_view_item(self, attr: str):
        """
        Return the first class in the MRO that defines `trait_name` in its __dict__.
        Works for Traits since class attributes are descriptors stored on the class.
        """
        if not attr[0].islower():
            return False

        here = False
        for cls in type(self).__mro__:
            if hasattr(cls, 'class_traits'):
                somewhere = attr in cls.class_traits()
                if cls.__module__ in ('__main__', 'cv'):
                    here |= somewhere
                else:
                    here &= not somewhere

        return here


proposal = Proposal(file='cover_and_cv.md')


def propose(*files, edit=False):
    global proposal

    try:
        from pyface.api import GUI
    except ImportError:
        raise ImportError("pip install pyside6<6.6  # Only available in Python <3.11")

    for file in files:
        # Import over existing content, keeping original order
        proposal.file = file

    proposal.edit_traits(view=proposal.view())
    GUI().start_event_loop()

def main():
    import fire
    fire.Fire(propose)

    from build_cover_letter_and_cv_markdown import MarkdownBuilder

    cv = f"{proposal.name}_CV.docx"
    cover = f"{proposal.name}-{proposal.role}-{proposal.organization}.docx"
    proposal.build_cv(cv)
    proposal.build_cover(cover)
    MarkdownBuilder(cover=cover, cv=cv)()


if __name__ == "__main__":
    main()
